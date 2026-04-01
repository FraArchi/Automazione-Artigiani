import json
import os
import smtplib
from datetime import datetime, timezone
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy import Column, DateTime, Float, ForeignKey, Integer, String, Text, create_engine
from sqlalchemy.orm import declarative_base, relationship, sessionmaker

print("🚀 Avvio applicazione...")

app = FastAPI(title="Automazione Artigiani")
Base = declarative_base()

ENVIRONMENT = os.environ.get("ENVIRONMENT", "development")
DATABASE_URL = os.environ.get("DATABASE_URL", "sqlite:///./local.db")
DEFAULT_RECEIVER_EMAIL = os.environ.get("DEFAULT_RECEIVER_EMAIL", "")
SENDER_EMAIL = os.environ.get("SENDER_EMAIL", "")
SENDER_PASSWORD = os.environ.get("SENDER_PASSWORD", "")
PUBLIC_BASE_URL = os.environ.get("PUBLIC_BASE_URL", "")
QUOTE_OUTPUT_DIR = Path(os.environ.get("QUOTE_OUTPUT_DIR", "quotes"))
QUOTE_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


print(f"📡 DATABASE_URL trovato: {'Sì' if DATABASE_URL else 'No'}")

if DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

engine_kwargs: dict[str, Any] = {}
if DATABASE_URL.startswith("sqlite"):
    engine_kwargs["connect_args"] = {"check_same_thread": False}

engine = create_engine(DATABASE_URL, **engine_kwargs)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


class Config(Base):
    __tablename__ = "config"

    id = Column(Integer, primary_key=True, index=True)
    key = Column(String, unique=True, index=True, nullable=False)
    value = Column(String, nullable=False)


class Lead(Base):
    __tablename__ = "leads"

    id = Column(Integer, primary_key=True, index=True)
    source = Column(String, default="webhook", nullable=False)
    client_name = Column(String, nullable=True)
    client_email = Column(String, nullable=True)
    client_phone = Column(String, nullable=True)
    job_type = Column(String, nullable=True)
    description = Column(Text, nullable=True)
    raw_payload = Column(Text, nullable=False)
    normalized_payload = Column(Text, nullable=False)
    status = Column(String, default="new", nullable=False)
    missing_fields = Column(Text, default="[]", nullable=False)
    review_summary = Column(Text, nullable=True)
    suggested_action = Column(String, nullable=True)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), nullable=False)
    updated_at = Column(
        DateTime,
        default=lambda: datetime.now(timezone.utc),
        onupdate=lambda: datetime.now(timezone.utc),
        nullable=False,
    )

    quotes = relationship("Quote", back_populates="lead", cascade="all, delete-orphan")


class Quote(Base):
    __tablename__ = "quotes"

    id = Column(Integer, primary_key=True, index=True)
    lead_id = Column(Integer, ForeignKey("leads.id"), nullable=False, index=True)
    file_path = Column(String, nullable=False)
    subtotal = Column(Float, nullable=False)
    vat = Column(Float, nullable=False)
    total = Column(Float, nullable=False)
    version = Column(Integer, default=1, nullable=False)
    status = Column(String, default="draft", nullable=False)
    generated_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), nullable=False)
    sent_at = Column(DateTime, nullable=True)

    lead = relationship("Lead", back_populates="quotes")


class ActivityLog(Base):
    __tablename__ = "activity_log"

    id = Column(Integer, primary_key=True, index=True)
    lead_id = Column(Integer, ForeignKey("leads.id"), nullable=True, index=True)
    event_type = Column(String, nullable=False)
    message = Column(Text, nullable=False)
    actor = Column(String, default="system", nullable=False)
    created_at = Column(DateTime, default=lambda: datetime.now(timezone.utc), nullable=False)


Base.metadata.create_all(bind=engine)
print("✅ Database inizializzato correttamente.")

if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")
    print("📂 Cartella 'static' montata.")
else:
    print("⚠️ Cartella 'static' non trovata!")


def now_utc() -> datetime:
    return datetime.now(timezone.utc)


def env_bool(name: str, default: bool = False) -> bool:
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on"}


def dumps_json(data: Any) -> str:
    return json.dumps(data, ensure_ascii=False, sort_keys=True)


def loads_json(data: str | None, default: Any) -> Any:
    if not data:
        return default
    try:
        return json.loads(data)
    except json.JSONDecodeError:
        return default


def serialize_datetime(value: datetime | None) -> str | None:
    return value.isoformat() if value else None


def serialize_activity(item: ActivityLog) -> dict[str, Any]:
    return {
        "id": item.id,
        "lead_id": item.lead_id,
        "event_type": item.event_type,
        "message": item.message,
        "actor": item.actor,
        "created_at": serialize_datetime(item.created_at),
    }


def detect_source(request: Request, payload: dict[str, Any]) -> str:
    headers = {key.lower(): value for key, value in request.headers.items()}
    data_section = payload.get("data") if isinstance(payload, dict) else None

    if headers.get("x-tally-signature") or headers.get("x-tally-event-id"):
        return "tally_webhook"

    if isinstance(data_section, dict) and (
        data_section.get("formId")
        or data_section.get("responseId")
        or isinstance(data_section.get("fields"), list)
    ):
        return "tally_webhook"

    provider = (payload.get("source") or payload.get("provider") or "").strip().lower() if isinstance(payload, dict) else ""
    if provider:
        return provider.replace(" ", "_")

    return "webhook"


def find_duplicate_lead(db, source: str, payload: dict[str, Any]) -> Lead | None:
    raw_payload = dumps_json(payload)
    return (
        db.query(Lead)
        .filter(Lead.source == source, Lead.raw_payload == raw_payload)
        .order_by(Lead.id.desc())
        .first()
    )


def db_session():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def get_config(db, key: str, default: str = "") -> str:
    config = db.query(Config).filter(Config.key == key).first()
    return config.value if config else default


def set_config(db, key: str, value: str) -> None:
    config = db.query(Config).filter(Config.key == key).first()
    if config:
        config.value = value
    else:
        db.add(Config(key=key, value=value))
    db.commit()


def bootstrap_config() -> None:
    db = SessionLocal()
    try:
        if not get_config(db, "active_receiver_email", ""):
            fallback = DEFAULT_RECEIVER_EMAIL or SENDER_EMAIL or "owner@example.com"
            set_config(db, "active_receiver_email", fallback)
        if not get_config(db, "email_enabled", ""):
            set_config(db, "email_enabled", "true" if env_bool("EMAIL_ENABLED", False) else "false")
    finally:
        db.close()


bootstrap_config()


def log_event(db, event_type: str, message: str, lead_id: int | None = None, actor: str = "system") -> None:
    db.add(ActivityLog(lead_id=lead_id, event_type=event_type, message=message, actor=actor))
    db.commit()


def parse_currency(value: Any) -> float | None:
    if value is None:
        return None
    text = str(value).strip().replace("€", "")
    text = text.replace(" ", "")
    if "," in text and "." not in text:
        text = text.replace(",", ".")
    else:
        text = text.replace(",", "")
    if not text:
        return None
    try:
        return float(text)
    except ValueError:
        return None


def parse_number(value: Any) -> float | None:
    return parse_currency(value)


def normalize_payload(raw_payload: dict[str, Any]) -> dict[str, Any]:
    extracted = raw_payload
    if isinstance(raw_payload, dict) and isinstance(raw_payload.get("data"), dict):
        data_section = raw_payload.get("data", {})
        if isinstance(data_section.get("fields"), list):
            extracted = {field.get("label"): field.get("value") for field in data_section["fields"]}

    normalized = {
        "client_name": extracted.get("Cliente") or extracted.get("Nome") or extracted.get("Nome del Cliente") or extracted.get("client_name"),
        "client_email": extracted.get("Email") or extracted.get("email") or extracted.get("client_email"),
        "client_phone": extracted.get("Telefono") or extracted.get("Telefono/WhatsApp") or extracted.get("client_phone"),
        "client_address": extracted.get("Indirizzo del Cliente") or extracted.get("Indirizzo") or extracted.get("client_address"),
        "job_type": extracted.get("Mestiere") or extracted.get("Mestiere / tipo di attività") or extracted.get("Tipologia_Richiesta") or extracted.get("job_type"),
        "description": extracted.get("Lavoro") or extracted.get("Descrizione del lavoro da svolgere") or extracted.get("Dettagli") or extracted.get("description"),
        "ore": extracted.get("Ore") or extracted.get("Ore di lavoro stimate") or extracted.get("ore"),
        "costo_orario": extracted.get("Costo orario") or extracted.get("Costo orario (€)") or extracted.get("costo_orario"),
        "materiali": extracted.get("Prezzo materiali (€)") or extracted.get("Materiali") or extracted.get("materiali"),
        "materiali_descrizione": extracted.get("Materiali necessari") or extracted.get("materiali_descrizione"),
        "artisan_name": extracted.get("Nome artigiano") or extracted.get("artisan_name"),
        "artisan_email": extracted.get("Email artigiano") or extracted.get("artisan_email"),
        "notes": extracted.get("Note") or extracted.get("Eventuali note aggiuntive") or extracted.get("notes"),
        "urgency": extracted.get("Urgenza del lavoro") or extracted.get("urgency"),
    }
    return normalized


REQUIRED_FOR_DRAFT = {
    "client_name": "cliente",
    "description": "descrizione",
    "ore": "ore",
    "costo_orario": "costo_orario",
    "materiali": "materiali",
}


def analyze_lead(normalized: dict[str, Any]) -> tuple[str, list[str], str, str]:
    missing = [field for field in REQUIRED_FOR_DRAFT if not normalized.get(field)]
    if missing:
        summary = f"Richiesta incompleta: mancano {', '.join(missing)}."
        return "incomplete", missing, summary, "request_missing_fields"

    summary = "Richiesta pronta per bozza preventivo. Verifica i costi e approva l'invio finale."
    return "ready_for_quote", [], summary, "review_draft"


def calculate_costs(normalized: dict[str, Any]) -> dict[str, float]:
    ore = parse_number(normalized.get("ore"))
    costo_orario = parse_currency(normalized.get("costo_orario"))
    materiali = parse_currency(normalized.get("materiali"))

    if ore is None or costo_orario is None or materiali is None:
        raise ValueError("Dati insufficienti per calcolare il preventivo")

    manodopera = ore * costo_orario
    subtotale = manodopera + materiali
    vat = subtotale * 0.22
    totale = subtotale + vat

    return {
        "ore": ore,
        "costo_orario": costo_orario,
        "materiali": materiali,
        "manodopera": manodopera,
        "subtotale": subtotale,
        "vat": vat,
        "total": totale,
    }


def safe_filename(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in value.strip())
    return cleaned or "cliente"


def create_professional_quote(lead: Lead, normalized: dict[str, Any], costs: dict[str, float], version: int) -> Path:
    client_name = normalized.get("client_name") or "cliente"
    filename = f"Preventivo_{safe_filename(client_name)}_v{version}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = QUOTE_OUTPUT_DIR / filename

    doc = Document()
    artisan_name = normalized.get("artisan_name") or "Artigiano"
    job_type = normalized.get("job_type") or "Professionista"

    header = doc.add_paragraph()
    header.add_run(f"{artisan_name}\n").bold = True
    header.add_run(f"{job_type}\n")
    header.alignment = 2

    doc.add_heading("PREVENTIVO DI SPESA", 0)
    client_address = normalized.get("client_address")
    materials_description = normalized.get("materiali_descrizione")
    urgency = normalized.get("urgency")

    header_text = (
        f"Data: {datetime.now().strftime('%d/%m/%Y')}\n"
        f"Validità: 30 giorni\n"
        f"Cliente: {client_name}"
    )
    if client_address:
        header_text += f"\nIndirizzo: {client_address}"
    doc.add_paragraph(header_text)
    doc.add_heading(f"Spett.le {client_name}", level=2)
    doc.add_paragraph(f"Oggetto: {normalized.get('description', 'Intervento')}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Descrizione"
    headers[1].text = "Quantità"
    headers[2].text = "Prezzo Un."
    headers[3].text = "Totale"

    manodopera_row = table.add_row().cells
    manodopera_row[0].text = "Manodopera"
    manodopera_row[1].text = f"{costs['ore']:.2f} h"
    manodopera_row[2].text = f"€{costs['costo_orario']:.2f}"
    manodopera_row[3].text = f"€{costs['manodopera']:.2f}"

    materiali_row = table.add_row().cells
    materiali_row[0].text = f"Materiali{': ' + str(materials_description) if materials_description else ''}"
    materiali_row[1].text = "1"
    materiali_row[2].text = f"€{costs['materiali']:.2f}"
    materiali_row[3].text = f"€{costs['materiali']:.2f}"

    if urgency or normalized.get("notes"):
        details = doc.add_paragraph()
        if urgency:
            details.add_run(f"Urgenza: {urgency}\n")
        if normalized.get("notes"):
            details.add_run(f"Note: {normalized.get('notes')}")

    footer = doc.add_paragraph()
    footer.alignment = 2
    footer.add_run(f"\nSubtotale: €{costs['subtotale']:.2f}\n")
    footer.add_run(f"IVA 22%: €{costs['vat']:.2f}\n")
    footer.add_run(f"TOTALE: €{costs['total']:.2f}").bold = True

    doc.save(path)
    return path


def serialize_quote(quote: Quote) -> dict[str, Any]:
    return {
        "id": quote.id,
        "lead_id": quote.lead_id,
        "file_path": quote.file_path,
        "subtotal": quote.subtotal,
        "vat": quote.vat,
        "total": quote.total,
        "version": quote.version,
        "status": quote.status,
        "generated_at": serialize_datetime(quote.generated_at),
        "sent_at": serialize_datetime(quote.sent_at),
    }


def serialize_lead(lead: Lead) -> dict[str, Any]:
    latest_quote = None
    if lead.quotes:
        latest_quote = sorted(lead.quotes, key=lambda q: (q.version, q.id))[-1]

    return {
        "id": lead.id,
        "source": lead.source,
        "client_name": lead.client_name,
        "client_email": lead.client_email,
        "client_phone": lead.client_phone,
        "job_type": lead.job_type,
        "description": lead.description,
        "status": lead.status,
        "missing_fields": loads_json(lead.missing_fields, []),
        "review_summary": lead.review_summary,
        "suggested_action": lead.suggested_action,
        "normalized_payload": loads_json(lead.normalized_payload, {}),
        "created_at": serialize_datetime(lead.created_at),
        "updated_at": serialize_datetime(lead.updated_at),
        "latest_quote": serialize_quote(latest_quote) if latest_quote else None,
    }


def get_active_receiver(db) -> str:
    return get_config(db, "active_receiver_email", DEFAULT_RECEIVER_EMAIL or SENDER_EMAIL or "owner@example.com")


def set_active_receiver(db, email: str) -> None:
    set_config(db, "active_receiver_email", email)


def email_sending_enabled(db) -> bool:
    config_value = get_config(db, "email_enabled", "true" if env_bool("EMAIL_ENABLED", False) else "false")
    return config_value.lower() == "true"


def build_public_url(path: str) -> str:
    base_url = PUBLIC_BASE_URL.rstrip("/")
    if not base_url:
        return path
    return f"{base_url}{path}"


def send_email_message(receiver: str, subject: str, body: str, attachment_path: str | None = None) -> None:
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = SENDER_EMAIL
    msg["To"] = receiver
    msg.attach(MIMEText(body))

    if attachment_path:
        with open(attachment_path, "rb") as file_handle:
            attachment_name = Path(attachment_path).name
            attachment = MIMEApplication(file_handle.read(), Name=attachment_name)
            attachment["Content-Disposition"] = f'attachment; filename="{attachment_name}"'
            msg.attach(attachment)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.send_message(msg)


def notify_artisan_quote_ready(db, lead: Lead, quote: Quote) -> bool:
    if not email_sending_enabled(db) or not SENDER_EMAIL or not SENDER_PASSWORD:
        return False

    normalized = loads_json(lead.normalized_payload, {})
    artisan_email = (normalized.get("artisan_email") or "").strip()
    if not artisan_email:
        return False

    download_url = build_public_url(f"/api/quotes/{quote.id}/download")
    body = (
        f"Ciao {normalized.get('artisan_name') or 'artigiano'},\n\n"
        f"La bozza del preventivo per {lead.client_name or 'cliente'} è pronta.\n"
        f"Descrizione: {lead.description or 'n/d'}\n"
        f"Totale bozza: €{quote.total:.2f}\n\n"
        f"Puoi scaricare il documento qui: {download_url}\n"
        f"In allegato trovi anche il file .docx.\n\n"
        "Controlla il preventivo prima di inviarlo al cliente finale."
    )
    send_email_message(
        artisan_email,
        f"Bozza preventivo pronta: {lead.client_name or 'cliente'}",
        body,
        attachment_path=quote.file_path,
    )
    log_event(db, "quote_ready_notified", f"Bozza notificata all'artigiano {artisan_email}", lead_id=lead.id, actor="system")
    return True


def generate_quote_for_lead(db, lead: Lead) -> Quote:
    normalized = loads_json(lead.normalized_payload, {})
    costs = calculate_costs(normalized)
    current_version = max((quote.version for quote in lead.quotes), default=0) + 1
    file_path = create_professional_quote(lead, normalized, costs, current_version)

    quote = Quote(
        lead_id=lead.id,
        file_path=str(file_path),
        subtotal=costs["subtotale"],
        vat=costs["vat"],
        total=costs["total"],
        version=current_version,
        status="draft",
    )
    db.add(quote)
    lead.status = "ready_for_quote"
    lead.updated_at = now_utc()
    db.commit()
    db.refresh(quote)
    log_event(db, "quote_generated", f"Bozza preventivo v{current_version} generata", lead_id=lead.id)
    try:
        notify_artisan_quote_ready(db, lead, quote)
    except Exception as exc:
        log_event(db, "quote_notification_failed", f"Notifica artigiano fallita: {exc}", lead_id=lead.id, actor="system")
    return quote


def send_quote_email(db, quote: Quote) -> None:
    if not email_sending_enabled(db):
        raise HTTPException(status_code=409, detail="Email sending is disabled in this environment")

    if not SENDER_EMAIL or not SENDER_PASSWORD:
        raise HTTPException(status_code=409, detail="Email credentials are not configured")

    lead = db.query(Lead).filter(Lead.id == quote.lead_id).first()
    if not lead:
        raise HTTPException(status_code=404, detail="Lead not found for quote")

    receiver = get_active_receiver(db)
    send_email_message(
        receiver,
        f"Nuovo Preventivo: {Path(quote.file_path).name}",
        "Ciao,\n\nIn allegato trovi il preventivo generato dal sistema.\n\nRicorda di verificare i dettagli prima dell'invio definitivo al cliente.",
        attachment_path=quote.file_path,
    )

    quote.status = "sent"
    quote.sent_at = now_utc()
    lead.status = "sent"
    lead.updated_at = now_utc()
    db.commit()
    log_event(db, "quote_sent", f"Preventivo inviato a {receiver}", lead_id=lead.id)


@app.get("/")
async def serve_dashboard():
    if os.path.exists("static/index.html"):
        return FileResponse("static/index.html")
    return {"message": "Dashboard non trovata"}


@app.get("/api/dashboard/summary")
async def api_dashboard_summary():
    db = SessionLocal()
    try:
        leads = db.query(Lead).all()
        recent_activity = db.query(ActivityLog).order_by(ActivityLog.created_at.desc(), ActivityLog.id.desc()).limit(10).all()
        counts = {
            "new": 0,
            "incomplete": 0,
            "needs_review": 0,
            "ready_for_quote": 0,
            "quoted": 0,
            "sent": 0,
            "error": 0,
            "archived": 0,
            "draft_quotes": db.query(Quote).filter(Quote.status == "draft").count(),
        }
        for lead in leads:
            counts.setdefault(lead.status, 0)
            counts[lead.status] += 1

        return {
            "counts": counts,
            "active_receiver_email": get_active_receiver(db),
            "email_enabled": email_sending_enabled(db),
            "recent_activity": [serialize_activity(item) for item in recent_activity],
        }
    finally:
        db.close()


@app.get("/api/activity-log")
async def api_activity_log(limit: int = 20):
    db = SessionLocal()
    try:
        safe_limit = max(1, min(limit, 100))
        items = db.query(ActivityLog).order_by(ActivityLog.created_at.desc(), ActivityLog.id.desc()).limit(safe_limit).all()
        return {"items": [serialize_activity(item) for item in items]}
    finally:
        db.close()


@app.get("/api/leads")
async def api_get_leads():
    db = SessionLocal()
    try:
        leads = db.query(Lead).order_by(Lead.created_at.desc()).all()
        return {"leads": [serialize_lead(lead) for lead in leads]}
    finally:
        db.close()


@app.get("/api/leads/{lead_id}")
async def api_get_lead(lead_id: int):
    db = SessionLocal()
    try:
        lead = db.query(Lead).filter(Lead.id == lead_id).first()
        if not lead:
            raise HTTPException(status_code=404, detail="Lead not found")
        return serialize_lead(lead)
    finally:
        db.close()


@app.patch("/api/leads/{lead_id}/review")
async def api_update_lead_review(lead_id: int, request: Request):
    data = await request.json()
    db = SessionLocal()
    try:
        lead = db.query(Lead).filter(Lead.id == lead_id).first()
        if not lead:
            raise HTTPException(status_code=404, detail="Lead not found")

        if "status" in data:
            lead.status = data["status"]
        if "missing_fields" in data:
            lead.missing_fields = dumps_json(data["missing_fields"])
        if "review_summary" in data:
            lead.review_summary = data["review_summary"]
        if "suggested_action" in data:
            lead.suggested_action = data["suggested_action"]
        lead.updated_at = now_utc()
        db.commit()
        log_event(db, "lead_review_updated", "Revisione lead aggiornata", lead_id=lead.id, actor="hermes")
        db.refresh(lead)
        return serialize_lead(lead)
    finally:
        db.close()


@app.post("/api/leads/{lead_id}/generate-quote")
async def api_generate_quote(lead_id: int):
    db = SessionLocal()
    try:
        lead = db.query(Lead).filter(Lead.id == lead_id).first()
        if not lead:
            raise HTTPException(status_code=404, detail="Lead not found")
        if lead.status != "ready_for_quote":
            raise HTTPException(status_code=409, detail="Lead is not ready for quote generation")

        quote = generate_quote_for_lead(db, lead)
        return {"status": "success", "quote": serialize_quote(quote), "lead": serialize_lead(lead)}
    finally:
        db.close()


@app.get("/api/quotes/{quote_id}/download")
async def api_download_quote(quote_id: int):
    db = SessionLocal()
    try:
        quote = db.query(Quote).filter(Quote.id == quote_id).first()
        if not quote:
            raise HTTPException(status_code=404, detail="Quote not found")
        if not os.path.exists(quote.file_path):
            raise HTTPException(status_code=404, detail="Quote file not found")
        return FileResponse(
            quote.file_path,
            filename=Path(quote.file_path).name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    finally:
        db.close()


@app.post("/api/quotes/{quote_id}/send")
async def api_send_quote(quote_id: int):
    db = SessionLocal()
    try:
        quote = db.query(Quote).filter(Quote.id == quote_id).first()
        if not quote:
            raise HTTPException(status_code=404, detail="Quote not found")
        send_quote_email(db, quote)
        db.refresh(quote)
        return {"status": "success", "quote": serialize_quote(quote)}
    finally:
        db.close()


@app.get("/api/receivers")
async def api_get_receivers():
    db = SessionLocal()
    try:
        return {"receivers": [get_active_receiver(db)]}
    finally:
        db.close()


@app.post("/api/receivers")
async def api_set_receiver(request: Request):
    data = await request.json()
    email = (data.get("email") or "").strip()
    if not email:
        raise HTTPException(status_code=400, detail="Email mancante")

    db = SessionLocal()
    try:
        set_active_receiver(db, email)
        log_event(db, "receiver_updated", f"Destinatario attivo impostato a {email}")
        return {"status": "success", "receiver": email}
    finally:
        db.close()


@app.delete("/api/receivers/{email}")
async def api_delete_receiver(email: str):
    db = SessionLocal()
    try:
        fallback = DEFAULT_RECEIVER_EMAIL or SENDER_EMAIL or "owner@example.com"
        set_active_receiver(db, fallback)
        log_event(db, "receiver_reset", f"Destinatario resettato da {email} a {fallback}")
        return {"status": "success", "receiver": fallback}
    finally:
        db.close()


@app.get("/webhook")
async def webhook_info():
    return {
        "message": "Webhook endpoint attivo. Usa POST per inviare i payload del form.",
        "methods": ["GET", "POST", "OPTIONS"],
    }


@app.options("/webhook", status_code=204)
async def webhook_options():
    return None


@app.post("/webhook")
async def receive_form(request: Request):
    try:
        payload = await request.json()
    except Exception:
        form_data = await request.form()
        payload = dict(form_data)

    source = detect_source(request, payload)
    normalized = normalize_payload(payload)
    initial_status = "new"
    initial_missing_fields: list[str] = []
    initial_review_summary = "Lead acquisito. In attesa della revisione Hermes."
    initial_suggested_action = "run_reviewer"

    db = SessionLocal()
    try:
        duplicate_lead = find_duplicate_lead(db, source, payload)
        if duplicate_lead:
            latest_quote = sorted(duplicate_lead.quotes, key=lambda quote: (quote.version, quote.id))[-1] if duplicate_lead.quotes else None
            log_event(
                db,
                "webhook_duplicate",
                f"Webhook duplicato ignorato per {duplicate_lead.client_name or 'cliente sconosciuto'}",
                lead_id=duplicate_lead.id,
            )
            return {
                "status": "duplicate",
                "duplicate": True,
                "lead_id": duplicate_lead.id,
                "lead_status": duplicate_lead.status,
                "quote_id": latest_quote.id if latest_quote else None,
                "quote_status": latest_quote.status if latest_quote else None,
                "missing_fields": loads_json(duplicate_lead.missing_fields, []),
                "review_summary": duplicate_lead.review_summary,
                "suggested_action": duplicate_lead.suggested_action,
            }

        lead = Lead(
            source=source,
            client_name=normalized.get("client_name"),
            client_email=normalized.get("client_email"),
            client_phone=normalized.get("client_phone"),
            job_type=normalized.get("job_type"),
            description=normalized.get("description"),
            raw_payload=dumps_json(payload),
            normalized_payload=dumps_json(normalized),
            status=initial_status,
            missing_fields=dumps_json(initial_missing_fields),
            review_summary=initial_review_summary,
            suggested_action=initial_suggested_action,
        )
        db.add(lead)
        db.commit()
        db.refresh(lead)
        log_event(db, "lead_created", f"Nuova richiesta ricevuta per {lead.client_name or 'cliente sconosciuto'}", lead_id=lead.id)

        return {
            "status": "success",
            "lead_id": lead.id,
            "lead_status": lead.status,
            "quote_id": None,
            "quote_status": None,
            "missing_fields": initial_missing_fields,
            "review_summary": initial_review_summary,
            "suggested_action": initial_suggested_action,
        }
    finally:
        db.close()


print("🚀 Applicazione pronta e in ascolto.")
