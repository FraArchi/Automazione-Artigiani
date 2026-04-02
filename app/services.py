from __future__ import annotations

from datetime import datetime
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import HTTPException, Request

from app.models import ActivityLog, Config, Lead, Quote
from app.settings import Settings
from app.utils import dumps_json, loads_json, now_utc, parse_currency, parse_number, safe_filename, serialize_datetime

REQUIRED_FOR_DRAFT = {
    "client_name": "cliente",
    "description": "descrizione",
    "ore": "ore",
    "costo_orario": "costo_orario",
    "materiali": "materiali",
}

FIELD_MAPPING = {
    "client_name": ["Cliente", "Nome", "Nome del Cliente", "client_name"],
    "client_email": ["Email", "email", "client_email"],
    "client_phone": ["Telefono", "Telefono/WhatsApp", "client_phone"],
    "client_address": ["Indirizzo del Cliente", "Indirizzo", "client_address"],
    "job_type": ["Mestiere", "Mestiere / tipo di attività", "Tipologia_Richiesta", "job_type"],
    "description": ["Lavoro", "Descrizione del lavoro da svolgere", "Dettagli", "description"],
    "ore": ["Ore", "Ore di lavoro stimate", "ore"],
    "costo_orario": ["Costo orario", "Costo orario (€)", "costo_orario"],
    "materiali": ["Prezzo materiali (€)", "Materiali", "materiali"],
    "materiali_descrizione": ["Materiali necessari", "materiali_descrizione"],
    "artisan_name": ["Nome artigiano", "artisan_name"],
    "artisan_email": ["Email artigiano", "artisan_email"],
    "notes": ["Note", "Eventuali note aggiuntive", "notes"],
    "urgency": ["Urgenza del lavoro", "urgency"],
}


def serialize_activity(item: ActivityLog) -> dict[str, Any]:
    return {
        "id": item.id,
        "lead_id": item.lead_id,
        "event_type": item.event_type,
        "message": item.message,
        "actor": item.actor,
        "created_at": serialize_datetime(item.created_at),
    }


def serialize_quote(quote: Quote) -> dict[str, Any]:
    return {
        "id": quote.id,
        "lead_id": quote.lead_id,
        "file_path": quote.file_path,
        "subtotal": quote.subtotal,
        "vat": quote.vat,
        "total": quote.total,
        "version": quote.version,
        "status": quote.status,
        "generated_at": serialize_datetime(quote.generated_at),
        "sent_at": serialize_datetime(quote.sent_at),
    }


def serialize_lead(lead: Lead) -> dict[str, Any]:
    latest_quote = None
    if lead.quotes:
        latest_quote = sorted(lead.quotes, key=lambda quote: (quote.version, quote.id))[-1]

    return {
        "id": lead.id,
        "source": lead.source,
        "client_name": lead.client_name,
        "client_email": lead.client_email,
        "client_phone": lead.client_phone,
        "job_type": lead.job_type,
        "description": lead.description,
        "status": lead.status,
        "missing_fields": loads_json(lead.missing_fields, []),
        "review_summary": lead.review_summary,
        "suggested_action": lead.suggested_action,
        "normalized_payload": loads_json(lead.normalized_payload, {}),
        "created_at": serialize_datetime(lead.created_at),
        "updated_at": serialize_datetime(lead.updated_at),
        "latest_quote": serialize_quote(latest_quote) if latest_quote else None,
    }


def detect_source(request: Request, payload: dict[str, Any]) -> str:
    headers = {key.lower(): value for key, value in request.headers.items()}
    data_section = payload.get("data") if isinstance(payload, dict) else None

    if headers.get("x-tally-signature") or headers.get("x-tally-event-id"):
        return "tally_webhook"

    if isinstance(data_section, dict) and (
        data_section.get("formId")
        or data_section.get("responseId")
        or isinstance(data_section.get("fields"), list)
    ):
        return "tally_webhook"

    provider = (payload.get("source") or payload.get("provider") or "").strip().lower() if isinstance(payload, dict) else ""
    if provider:
        return provider.replace(" ", "_")

    return "webhook"


def find_duplicate_lead(db, source: str, payload: dict[str, Any]) -> Lead | None:
    raw_payload = dumps_json(payload)
    return (
        db.query(Lead)
        .filter(Lead.source == source, Lead.raw_payload == raw_payload)
        .order_by(Lead.id.desc())
        .first()
    )


def extract_payload_fields(raw_payload: dict[str, Any]) -> dict[str, Any]:
    extracted = raw_payload if isinstance(raw_payload, dict) else {}
    if isinstance(raw_payload, dict) and isinstance(raw_payload.get("data"), dict):
        data_section = raw_payload.get("data", {})
        if isinstance(data_section.get("fields"), list):
            extracted = {
                field.get("label"): field.get("value")
                for field in data_section["fields"]
                if isinstance(field, dict) and field.get("label")
            }
    return extracted if isinstance(extracted, dict) else {}


def normalize_payload(raw_payload: dict[str, Any]) -> dict[str, Any]:
    extracted = extract_payload_fields(raw_payload)

    normalized: dict[str, Any] = {}
    for target_field, source_labels in FIELD_MAPPING.items():
        normalized[target_field] = next(
            (extracted.get(label) for label in source_labels if extracted.get(label) not in (None, "")),
            None,
        )
    return normalized


def build_mapping_debug(raw_payload: dict[str, Any], normalized_payload: dict[str, Any]) -> dict[str, Any]:
    extracted = extract_payload_fields(raw_payload)
    mapped_source_labels: list[str] = []
    mapped_labels_by_field: dict[str, str] = {}

    for target_field, source_labels in FIELD_MAPPING.items():
        target_value = normalized_payload.get(target_field)
        if target_value in (None, ""):
            continue
        for label in source_labels:
            if extracted.get(label) == target_value:
                mapped_source_labels.append(label)
                mapped_labels_by_field[target_field] = label
                break

    unmapped_fields = {
        key: value
        for key, value in extracted.items()
        if key not in mapped_source_labels and key not in {"eventId", "eventType", "source", "provider"}
    }

    missing_critical_fields = [field for field in REQUIRED_FOR_DRAFT if normalized_payload.get(field) in (None, "")]
    if not normalized_payload.get("client_email") and not normalized_payload.get("client_phone"):
        missing_critical_fields.append("contact")

    return {
        "raw_payload": raw_payload,
        "extracted_fields": extracted,
        "normalized_payload": normalized_payload,
        "mapped_labels_by_field": mapped_labels_by_field,
        "unmapped_fields": unmapped_fields,
        "missing_critical_fields": missing_critical_fields,
    }


def analyze_lead(normalized: dict[str, Any]) -> tuple[str, list[str], str, str]:
    missing = [field for field in REQUIRED_FOR_DRAFT if not normalized.get(field)]
    if missing:
        summary = f"Richiesta incompleta: mancano {', '.join(missing)}."
        return "incomplete", missing, summary, "request_missing_fields"

    summary = "Richiesta pronta per bozza preventivo. Verifica i costi e approva l'invio finale."
    return "ready_for_quote", [], summary, "review_draft"


def calculate_costs(normalized: dict[str, Any]) -> dict[str, float]:
    ore = parse_number(normalized.get("ore"))
    costo_orario = parse_currency(normalized.get("costo_orario"))
    materiali = parse_currency(normalized.get("materiali"))

    if ore is None or costo_orario is None or materiali is None:
        raise ValueError("Dati insufficienti per calcolare il preventivo")

    manodopera = ore * costo_orario
    subtotale = manodopera + materiali
    vat = subtotale * 0.22
    totale = subtotale + vat

    return {
        "ore": ore,
        "costo_orario": costo_orario,
        "materiali": materiali,
        "manodopera": manodopera,
        "subtotale": subtotale,
        "vat": vat,
        "total": totale,
    }


def create_professional_quote(settings: Settings, lead: Lead, normalized: dict[str, Any], costs: dict[str, float], version: int) -> Path:
    client_name = normalized.get("client_name") or "cliente"
    filename = f"Preventivo_{safe_filename(client_name)}_v{version}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = settings.quote_output_dir / filename

    doc = Document()
    artisan_name = normalized.get("artisan_name") or "Artigiano"
    job_type = normalized.get("job_type") or "Professionista"

    header = doc.add_paragraph()
    header.add_run(f"{artisan_name}\n").bold = True
    header.add_run(f"{job_type}\n")
    header.alignment = 2

    doc.add_heading("PREVENTIVO DI SPESA", 0)
    client_address = normalized.get("client_address")
    materials_description = normalized.get("materiali_descrizione")
    urgency = normalized.get("urgency")

    header_text = (
        f"Data: {datetime.now().strftime('%d/%m/%Y')}\n"
        f"Validità: 30 giorni\n"
        f"Cliente: {client_name}"
    )
    if client_address:
        header_text += f"\nIndirizzo: {client_address}"
    doc.add_paragraph(header_text)
    doc.add_heading(f"Spett.le {client_name}", level=2)
    doc.add_paragraph(f"Oggetto: {normalized.get('description', 'Intervento')}")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Descrizione"
    headers[1].text = "Quantità"
    headers[2].text = "Prezzo Un."
    headers[3].text = "Totale"

    manodopera_row = table.add_row().cells
    manodopera_row[0].text = "Manodopera"
    manodopera_row[1].text = f"{costs['ore']:.2f} h"
    manodopera_row[2].text = f"€{costs['costo_orario']:.2f}"
    manodopera_row[3].text = f"€{costs['manodopera']:.2f}"

    materiali_row = table.add_row().cells
    materiali_row[0].text = f"Materiali{': ' + str(materials_description) if materials_description else ''}"
    materiali_row[1].text = "1"
    materiali_row[2].text = f"€{costs['materiali']:.2f}"
    materiali_row[3].text = f"€{costs['materiali']:.2f}"

    if urgency or normalized.get("notes"):
        details = doc.add_paragraph()
        if urgency:
            details.add_run(f"Urgenza: {urgency}\n")
        if normalized.get("notes"):
            details.add_run(f"Note: {normalized.get('notes')}")

    footer = doc.add_paragraph()
    footer.alignment = 2
    footer.add_run(f"\nSubtotale: €{costs['subtotale']:.2f}\n")
    footer.add_run(f"IVA 22%: €{costs['vat']:.2f}\n")
    footer.add_run(f"TOTALE: €{costs['total']:.2f}").bold = True

    doc.save(path)
    return path


def get_config(db, key: str, default: str = "") -> str:
    config = db.query(Config).filter(Config.key == key).first()
    return config.value if config else default


def set_config(db, key: str, value: str) -> None:
    config = db.query(Config).filter(Config.key == key).first()
    if config:
        config.value = value
    else:
        db.add(Config(key=key, value=value))
    db.commit()


def bootstrap_config(session_factory, settings: Settings) -> None:
    db = session_factory()
    try:
        if not get_config(db, "active_receiver_email", ""):
            fallback = settings.default_receiver_email or settings.sender_email or "owner@example.com"
            set_config(db, "active_receiver_email", fallback)
        if not get_config(db, "email_enabled", ""):
            set_config(db, "email_enabled", "true" if settings.email_enabled_default else "false")
    finally:
        db.close()


def log_event(db, event_type: str, message: str, lead_id: int | None = None, actor: str = "system") -> None:
    db.add(ActivityLog(lead_id=lead_id, event_type=event_type, message=message, actor=actor))
    db.commit()


def get_active_receiver(db, settings: Settings) -> str:
    return get_config(db, "active_receiver_email", settings.default_receiver_email or settings.sender_email or "owner@example.com")


def set_active_receiver(db, email: str) -> None:
    set_config(db, "active_receiver_email", email)


def email_sending_enabled(db, settings: Settings) -> bool:
    config_value = get_config(db, "email_enabled", "true" if settings.email_enabled_default else "false")
    return config_value.lower() == "true"


def build_public_url(settings: Settings, path: str) -> str:
    if not settings.public_base_url:
        return path
    return f"{settings.public_base_url}{path}"


def send_email_message(settings: Settings, smtplib_module, receiver: str, subject: str, body: str, attachment_path: str | None = None) -> None:
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = settings.sender_email
    msg["To"] = receiver
    msg.attach(MIMEText(body, _charset="utf-8"))

    if attachment_path:
        with open(attachment_path, "rb") as file_handle:
            attachment_name = Path(attachment_path).name
            attachment = MIMEApplication(file_handle.read(), Name=attachment_name)
            attachment["Content-Disposition"] = f'attachment; filename="{attachment_name}"'
            msg.attach(attachment)

    with smtplib_module.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(settings.sender_email, settings.sender_password)
        server.send_message(msg)


def notify_artisan_quote_ready(db, settings: Settings, smtplib_module, lead: Lead, quote: Quote) -> bool:
    if not email_sending_enabled(db, settings) or not settings.sender_email or not settings.sender_password:
        return False

    normalized = loads_json(lead.normalized_payload, {})
    artisan_email = (normalized.get("artisan_email") or "").strip()
    if not artisan_email:
        return False

    download_url = build_public_url(settings, f"/api/quotes/{quote.id}/download")
    body = (
        f"Ciao {normalized.get('artisan_name') or 'artigiano'},\n\n"
        f"La bozza del preventivo per {lead.client_name or 'cliente'} è pronta.\n"
        f"Descrizione: {lead.description or 'n/d'}\n"
        f"Totale bozza: €{quote.total:.2f}\n\n"
        f"Puoi scaricare il documento qui: {download_url}\n"
        f"In allegato trovi anche il file .docx.\n\n"
        "Controlla il preventivo prima di inviarlo al cliente finale."
    )
    send_email_message(
        settings,
        smtplib_module,
        artisan_email,
        f"Bozza preventivo pronta: {lead.client_name or 'cliente'}",
        body,
        attachment_path=quote.file_path,
    )
    log_event(db, "quote_ready_notified", f"Bozza notificata all'artigiano {artisan_email}", lead_id=lead.id, actor="system")
    return True


def generate_quote_for_lead(db, settings: Settings, smtplib_module, lead: Lead) -> Quote:
    normalized = loads_json(lead.normalized_payload, {})
    costs = calculate_costs(normalized)
    current_version = max((quote.version for quote in lead.quotes), default=0) + 1
    file_path = create_professional_quote(settings, lead, normalized, costs, current_version)

    quote = Quote(
        lead_id=lead.id,
        file_path=str(file_path),
        subtotal=costs["subtotale"],
        vat=costs["vat"],
        total=costs["total"],
        version=current_version,
        status="draft",
    )
    db.add(quote)
    lead.status = "ready_for_quote"
    lead.updated_at = now_utc()
    db.commit()
    db.refresh(quote)
    log_event(db, "quote_generated", f"Bozza preventivo v{current_version} generata", lead_id=lead.id)
    try:
        notify_artisan_quote_ready(db, settings, smtplib_module, lead, quote)
    except Exception as exc:
        log_event(db, "quote_notification_failed", f"Notifica artigiano fallita: {exc}", lead_id=lead.id, actor="system")
    return quote


def send_quote_email(db, settings: Settings, smtplib_module, quote: Quote) -> None:
    if not email_sending_enabled(db, settings):
        raise HTTPException(status_code=409, detail="Email sending is disabled in this environment")

    if not settings.sender_email or not settings.sender_password:
        raise HTTPException(status_code=409, detail="Email credentials are not configured")

    lead = db.query(Lead).filter(Lead.id == quote.lead_id).first()
    if not lead:
        raise HTTPException(status_code=404, detail="Lead not found for quote")

    receiver = get_active_receiver(db, settings)
    send_email_message(
        settings,
        smtplib_module,
        receiver,
        f"Nuovo Preventivo: {Path(quote.file_path).name}",
        "Ciao,\n\nIn allegato trovi il preventivo generato dal sistema.\n\nRicorda di verificare i dettagli prima dell'invio definitivo al cliente.",
        attachment_path=quote.file_path,
    )

    quote.status = "sent"
    quote.sent_at = now_utc()
    lead.status = "sent"
    lead.updated_at = now_utc()
    db.commit()
    log_event(db, "quote_sent", f"Preventivo inviato a {receiver}", lead_id=lead.id)
